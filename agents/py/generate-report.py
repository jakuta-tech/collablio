#!/usr/bin/env python3
#import http.server
#import socketserver
import json
import docx
from docx.shared import Pt
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_BREAK

import random
import sys
import re
import os
import time

import copy
import datetime
import uuid
import urllib.request
import urllib.parse
import traceback

import io
import mimetypes

######################################################
'''
Notes:
* There needs to be a ReportTemplate.docx in the current folder, 
* custom styles should be saved into the docx, along with headers, footers, a title page
* the auto generated content will be appended to anything that is in the template

TODO:
make this a microservice/daemon that polls the database for report nodes,
    if a report node has an annotation (eg. 'generate') then do the auto generation, 
    upload the report as a file attachment to the report node
    delete the annotation (create a folder called 'completed jobs' and move it there or maybe just the recycle bin)
'''
######################################################

C_USAGE = '''Usage: {0} <projectname> <reportname>
        generate a docx report
'''

COLLABLIO_HOST = 'http://127.0.0.1:5000'

PROP_UID = "uid"
PROP_TYPE = "ty"
PROP_LABEL = "l"
PROP_DETAIL = "d"
PROP_TEXTDATA = "x"
PROP_CUSTOM = "c"
PROP_TIME = "t"
PROP_LASTMOD = "m"
PROP_BINARYDATA = "b"
PROP_EDITING = "e"
PROP_PARENTLIST = "in"
PROP_CHILDLIST = "out"
PROP_RELATIONS = "lnk"

TYPE_CLIENT = "Client"
TYPE_PROJECT = "Project"
TYPE_FOLDER = "Folder"
TYPE_HOST = "Host"
TYPE_PORT = "Port"
TYPE_TEXT = "Text"
TYPE_IMAGE = "Image"
TYPE_FILE = "File"
TYPE_NOTE = "Note"
TYPE_TABLE = "Table"
TYPE_ANNOTATION = "Annotation"
TYPE_TAG = "Tag"
TYPE_REPORT = "Report"
TYPE_SECTION = "Section"
TYPE_JOBREQ = "Job Request"


DAEMON_LISTENER_MODE = False

def parseArgs():
    global DAEMON_LISTENER_MODE
    try:
        data = {}
        num_args = len(sys.argv)
        data['projectname'] = sys.argv[1]
        data['reportname'] = sys.argv[2]

        DAEMON_LISTENER_MODE = not (data['projectname'] and data['reportname'])
            #start in autogenerator listen mode


        return data
    except:
        print(C_USAGE.format(sys.argv[0]))
        exit(0)


#del foo.bar
class Node:
    
    def __init__(self, _type):
        self.Type = _type
        self.Children = []
        self.Parents = []
        self.UID = ''
        self.Label = ''
        self.Detail = ''
        self.CustomData = ''
        self.TextData = ''
        
    def convert(self):
        apiFormatNode = {}
        apiFormatNode[PROP_UID] = self.UID
        apiFormatNode[PROP_TYPE] = self.Type
        apiFormatNode[PROP_LABEL] = self.Label
        apiFormatNode[PROP_DETAIL] = self.Detail
        apiFormatNode[PROP_TEXTDATA] = self.TextData
        #only need to specify the parent UID for the host nodes if they don't have a UID (i.e. a new insert)
        apiFormatNode[PROP_PARENTLIST] = [{PROP_UID : self.Parents[0]}] if (len(self.Parents) > 0) else []
        apiFormatNode[PROP_CHILDLIST] = [{PROP_UID : child.UID} for child in self.Children]
        return apiFormatNode
        


def fetchNodes(query):
    response = urllib.request.urlopen(COLLABLIO_HOST+"/nodes"+query)
    jsonResponse =  json.loads(response.read().decode('utf8'))
    if 'nodes' not in jsonResponse:
        raise Exception()
    return jsonResponse


def recursiveConvertNodesToAPIFormat(node, listToAddTheNodeTo):
    listToAddTheNodeTo.append(node.convert())
    if node.Children:
        for child in node.Children:
            recursiveConvertNodesToAPIFormat(child, listToAddTheNodeTo)
        




class MultiPartForm:
    """Accumulate the data to be used when posting a form."""

    def __init__(self):
        self.form_fields = []
        self.files = []
        # Use a large random byte string to separate
        # parts of the MIME data.
        self.boundary = uuid.uuid4().hex.encode('utf-8')
        return

    def get_content_type(self):
        return 'multipart/form-data; boundary={}'.format(
            self.boundary.decode('utf-8'))

    def add_field(self, name, value):
        """Add a simple field to the form data."""
        self.form_fields.append((name, value))

    def add_file(self, fieldname, filename, fileHandle,
                 mimetype=None):
        """Add a file to be uploaded."""
        body = fileHandle.read()
        if mimetype is None:
            mimetype = (
                mimetypes.guess_type(filename)[0] or
                'application/octet-stream'
            )
        self.files.append((fieldname, filename, mimetype, body))
        return

    @staticmethod
    def _form_data(name):
        return ('Content-Disposition: form-data; '
                'name="{}"\r\n').format(name).encode('utf-8')

    @staticmethod
    def _attached_file(name, filename):
        #return ('Content-Disposition: file; '
        return ('Content-Disposition: form-data; '
                'name="{}"; filename="{}"\r\n').format(
                    name, filename).encode('utf-8')

    @staticmethod
    def _content_type(ct):
        return 'Content-Type: {}\r\n'.format(ct).encode('utf-8')

    def __bytes__(self):
        """Return a byte-string representing the form data,
        including attached files.
        """
        buffer = io.BytesIO()
        boundary = b'--' + self.boundary + b'\r\n'

        # Add the form fields
        for name, value in self.form_fields:
            buffer.write(boundary)
            buffer.write(self._form_data(name))
            buffer.write(b'\r\n')
            buffer.write(value.encode('utf-8'))
            buffer.write(b'\r\n')

        # Add the files to upload
        for f_name, filename, f_content_type, body in self.files:
            buffer.write(boundary)
            buffer.write(self._attached_file(f_name, filename))
            buffer.write(self._content_type(f_content_type))
            buffer.write(b'\r\n')
            buffer.write(body)
            buffer.write(b'\r\n')

        buffer.write(b'--' + self.boundary + b'--\r\n')
        return buffer.getvalue()
        


START_HEADING_LEVEL = 0
MAX_HEADING_LEVEL = 9

# https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
def applyStylesToRun(style, run):
    if not style:
        return
    if 'bold' in style:
        run.bold = True
    if 'italic' in style:
        run.italic = True


def initialiseDocStyles(doc):
    if 'Code Block' not in doc.styles:
        codeBlockStyle = doc.styles.add_style('Code Block',docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        codeBlockStyle.base_style = doc.styles['Body Text']
        codeBlockStyle.font.name = 'Consolas'
    
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def applyStylesToParagraph(style, paragraph, doc, startHeadingLevel):
    if not style:
        return
    docStyles = doc.styles
    pformat = paragraph.paragraph_format
    if 'code-block' in style:
        testTable = doc.add_table(1,1)
        testPara = testTable.cell(0,0).paragraphs[0]
        testTable.style = docStyles['Medium Shading 1 Accent 1'] #need to change this to use style Name (not style ID) to get rid of warning
        testTable.style.font.bold = False
        #testTable.style.font.italic = True
        testPara.style.font.bold = False
        testPara.paragraph_format.space_before = Cm(0.1)
        testPara.paragraph_format.line_spacing = 1

        for run in paragraph.runs:
            testPara.add_run(run.text, run.style)

        testPara.style = docStyles['Code Block']
        delete_paragraph(paragraph)
        trailingPara = doc.add_paragraph()
        trailingPara.paragraph_format.space_after = Cm(0)

        return # we've nuked paragraph, nothing else can reference it at this point

    if 'list' in style:
        listItemType = style['list']
        paragraph.style = docStyles['List Bullet']
        pformat.left_indent = Cm(1)
        pformat.space_before = Cm(0)
        pformat.space_after = Cm(0)

    if 'header' in style:
        headingLevel = int(style['header'])
        headingLevel += startHeadingLevel
        if headingLevel > MAX_HEADING_LEVEL:
            headingLevel = MAX_HEADING_LEVEL
        paragraph.style = docStyles['Heading '+str(headingLevel)]

    if 'blockquote' in style:
        paragraph.style = docStyles['Quote']


    #ensure this is last
    if 'indent' in style:
        pformat.left_indent += Cm(int(style['indent']))
        #run.text = '\t'*int(style['indent'])+run.text

STYLES_WITH_SEPARATE_PARAGRAPHS = ['list']

def isStyleWithItsOwnParagraph(style):
    if not style:
        return False
    for key in list(style.keys()):
        if key in STYLES_WITH_SEPARATE_PARAGRAPHS:
            return True
    #print('key '+key+' is not in STYLES_WITH_SEPARATE_PARAGRAPHS')
    return False


# formatting attributes in Quill are generally applied to the last operation's inserted line
# some attributes like bold, italic, strikethrough are applied to the line that follows (forward-applied attributes):
FWD_APPLIED_ATTRIBS = ['bold','italic']

# otherAttachments are things like tables, images to might be referenced in this quillOpsData
# its type is:  { "reference" : "unique_ref_string", "nodeUid":"uid_of_node" }
def convertQuillOpsToDOCX(quilljsondata, doc, curdepth, otherAttachments = None):

    quillOpsData = json.loads(quilljsondata)

    Lines = [{'content':'','formatting':''}]

    count = 1
    for op in quillOpsData['ops']:
        #print('\n\nLine #{} *********'.format(count))
        count = count+1
        #print(op)
        fwdAttrs = None
        if 'attributes' in op: 
            attrs = op['attributes']
            #print('attrs='+str(attrs))
            opHasFwdAppliedAttrib = False
            #print(list(attrs.keys()))
            for attrName in list(attrs.keys()):
                #print('checking FWD_APPLIED_ATTRIBS for '+attrName)
                opHasFwdAppliedAttrib = (opHasFwdAppliedAttrib or (attrName.lower() in FWD_APPLIED_ATTRIBS))
            if not opHasFwdAppliedAttrib:
                #print('No op attributes are in FWD_APPLY_ATTRIBS'+str(attrs))
                Lines[-1]['formatting'] = attrs
                #print('Lines[-1]: '+str(Lines[-1]))
            else:
                #print('Apply forward attribute'+str(attrs))
                fwdAttrs =  attrs

        startFromStrPos = 0
        while True:
            indexOfNextNewline = op['insert'].find('\n',startFromStrPos)
            if indexOfNextNewline < 0:
                #print('no more newlines found')
                Lines.append({ 'content': op['insert'][startFromStrPos:], 'formatting' : fwdAttrs })
                break
            #print('newline found at: '+str(indexOfNextNewline))
            Lines.append({ 'content': op['insert'][startFromStrPos:indexOfNextNewline+1], 'formatting' : fwdAttrs })
            startFromStrPos = indexOfNextNewline+1
            if startFromStrPos >= len(op['insert']):
                break

    cleanedLines = [Lines[0]]
    for iL in range(0, len(Lines)):
        line = Lines[iL]
        #print(line)
        #merge lines that are just {line:'\n','formatting':None} with previous line if previous line isn't a sole newline 
        lineBefore = Lines[iL-1]['content']
        if line['content'] == '\n' and not line['formatting'] and lineBefore != '\n':
            cleanedLines[-1]['content'] += '\n'
            #print('dropping empty line')
            continue
        else:
            cleanedLines.append(line)
            
    for line2 in cleanedLines:
        print(str(line2['formatting'])+' '+line2['content'])



    curLineIndex = 0
    prevLineStyle = {}
    runStr = ''
    prevStyleIsItsOwnParagraph = False

    while curLineIndex < len(cleanedLines):

        p = doc.add_paragraph('')
        finishedParagraph = False

        while not finishedParagraph:

            curLine = cleanedLines[curLineIndex]['content']
            curLineStyle = cleanedLines[curLineIndex]['formatting']

            prevStyleIsItsOwnParagraph = isStyleWithItsOwnParagraph(prevLineStyle)

            #print('*'*10+'\nprevStyle='+str(prevLineStyle)+' , curStyle='+str(curLineStyle)+', curLine=[['+curLine+']]')

            # the approach is to aggregate lines of the same style into the one paragraph
            # create a new paragraph once the style changes
            
            if ((curLineIndex >= len(cleanedLines) - 1) or ((curLineStyle != prevLineStyle) and runStr) or prevStyleIsItsOwnParagraph):
                #print('  - in add_run area')
                runEndsWithNewline = runStr.endswith('\n')
                if runEndsWithNewline:
                    runStr = runStr.strip('\n') #runStr[:-1] #get rid of trailing newline because space is automatically added between paragraphs
                run = p.add_run(runStr)
                
                if prevLineStyle:
                    applyStylesToRun(prevLineStyle, run)
                    
                #prevStyleIsItsOwnParagraph = isStyleWithItsOwnParagraph(prevLineStyle)
                #if prevStyleIsItsOwnParagraph:
                #    print('* '+str(prevLineStyle)+' is its own paragraph')
                finishedParagraph = (runEndsWithNewline or prevStyleIsItsOwnParagraph or (curLineIndex >= len(cleanedLines) - 1))
                if finishedParagraph:
                    applyStylesToParagraph(prevLineStyle, p, doc, curdepth)
                runStr = curLine #commence a new run string, assigning the current line
            else:
                runStr += curLine #append this line to the current run string

            prevLineStyle = curLineStyle
            curLineIndex += 1


def addAPageBreak(doc):
    doc.add_paragraph('').add_run('').add_break(WD_BREAK.PAGE)

def sortNodes(uidDict):
    return getNodeForUID(uidDict[PROP_UID])[PROP_LABEL]
    
# don't call this on the reportRootNode, ignore it and only create its child sections
def generate(docx, reportTreeNode, depth):
    print('generate')
    if not reportTreeNode:
        print('calling generate on a None reportTreeNode')
        return
    
    if reportTreeNode[PROP_TYPE] not in [TYPE_SECTION, TYPE_NOTE, TYPE_TEXT, TYPE_IMAGE, TYPE_TABLE]:
        return

    #assumes that the previous section has ended with a page break
    heading = 'Untitled Section'
    try:
        heading = reportTreeNode[PROP_LABEL]
    except:
        print('error when attempting to retrieve label for section '+reportTreeNode[PROP_UID])

    #putting "[N]" in the label is currently a dodgy hack to order the sections/notes, so remove it
    match = re.search('^\\[[a-z0-9]+\\][\\s]+', heading)
    if(match):
        heading = heading[len(match.group(0)):]
    #need to replace it with a proper x.y.z numbering scheme, by passing the parent section's number as a string to this function
    
    headingLevel = MAX_HEADING_LEVEL if (depth > MAX_HEADING_LEVEL) else depth

    p = docx.add_paragraph(heading, 'Heading '+str(headingLevel))

    #render table first??
    # or just require {{reference}} to tables and leave it up to user?
    if PROP_TEXTDATA in reportTreeNode:
        convertQuillOpsToDOCX(reportTreeNode[PROP_TEXTDATA], docx, depth)
        addAPageBreak(docx)

    if PROP_CHILDLIST in reportTreeNode:
        reportTreeNode[PROP_CHILDLIST].sort(key=sortNodes)
        for childSectionUID in reportTreeNode[PROP_CHILDLIST]:
            generate(docx, getNodeForUID(childSectionUID[PROP_UID]), depth+1)


NODE_INDEX = {'_':''}

def getNodeForUID(uid):
    global NODE_INDEX    
    node = NODE_INDEX[uid] if uid in NODE_INDEX else None
    return node

def storeNode(node):
    global NODE_INDEX
    NODE_INDEX[node[PROP_UID]] = node

def clearNodeIndex():
    global NODE_INDEX
    NODE_INDEX = {'_':''}
    
'''
public class QueryNodesPostData
{
    public List<string> uids {get; set;}
    public string field {get; set;}
    public string op {get; set;}
    public string val {get; set;}
    public int depth {get; set;}
    public string type {get; set;}
}
'''

def fetchNodesPost(uids = [], field = PROP_LASTMOD, op = 'gt', val = '0', depth = 20, typ = ''):

    reqdata = { 'uids': uids, 'field': field, 'op': op, 'val': val, 'depth': depth, 'type': typ }
    print(json.dumps(reqdata))
    req = urllib.request.Request(url=COLLABLIO_HOST+'/nodes', data=bytes(json.dumps(reqdata), encoding='utf-8'))
    req.add_header('Content-Type', 'application/json')
    response = urllib.request.urlopen(req)
    jsonResponse =  json.loads(response.read().decode('utf8'))
    if 'nodes' not in jsonResponse:
        raise Exception()
    return jsonResponse

def moveNodesPost(moveData):

    print(json.dumps(moveData))
    req = urllib.request.Request(url=COLLABLIO_HOST+'/move', data=bytes(json.dumps(moveData), encoding='utf-8'))
    req.add_header('Content-Type', 'application/json')
    response = urllib.request.urlopen(req)
    jsonResponse =  json.loads(response.read().decode('utf8'))
    return jsonResponse

def getRecycleBinFolderUID():
    querystring = '?uid={}&field={}&op={}&val={}&depth={}&type={}'.format(\
        '',\
        PROP_LABEL,\
        'eq',\
        urllib.parse.quote('Recycle Bin'),\
        20,\
        TYPE_FOLDER)
    print(COLLABLIO_HOST+"/nodes"+querystring)

    try:
        jsonResponse =  fetchNodes(querystring)
        return jsonResponse['nodes'][0][PROP_UID]
    except:
        print('recycle bin folder doesnt exist, creating a new one')

    try:
        newRecycleBinNode = Node(TYPE_FOLDER)
        newRecycleBinNode.Label = 'Recycle Bin'
        nodesToUpsert = []
        recursiveConvertNodesToAPIFormat(newRecycleBinNode, nodesToUpsert)
        serialisedJson = json.dumps(nodesToUpsert).encode('utf8')
        req = urllib.request.Request(COLLABLIO_HOST+'/upsert', data=serialisedJson, headers={'content-type': 'application/json'})
        response = urllib.request.urlopen(req)
        returnedUids = json.loads(response.read().decode('utf8'))
        return returnedUids[0]
    except Exception as e:
        print('an exception occurred while creating recycle bin: '+str(e))
        traceback.print_exc()


##################################################################################
## The main program
##################################################################################


argdata = parseArgs()

#QueryNodesGet(string uid = null, string field=null, string op=null, string val=null, int depth = 0, string type = null)
'''                
#query to find projectname data['projectname']
querystring = '?uid=&field={}&op={}&val={}&depth={}&type={}'.format(\
    PROP_LABEL,\
    'eq',\
    urllib.parse.quote(argdata['projectname']),\
    20,\
    TYPE_PROJECT)
print(COLLABLIO_HOST+"/nodes"+querystring)

try:
    jsonResponse =  fetchNodes(querystring)
    projectUID = ''
    for nodeResult in jsonResponse['nodes']:
        if nodeResult[PROP_LABEL] == argdata['projectname']:
            projectUID = nodeResult[PROP_UID]
            break
    print('located project '+projectUID)

    #query recursively to find the specified reportname under projectname
    querystring = '?uid={}&field={}&op={}&val={}&depth={}&type={}'.format(\
        projectUID,\
        PROP_LABEL,\
        'eq',\
        urllib.parse.quote(argdata['reportname']),\
        20,\
        TYPE_REPORT)
    print(COLLABLIO_HOST+"/nodes"+querystring)

    jsonResponse =  fetchNodes(querystring)

    reportRootNode = None
    for nodeResult in jsonResponse['nodes']:
        if nodeResult[PROP_LABEL] == argdata['reportname']:
            reportRootNode = nodeResult

    if not reportRootNode:
        print('unable to locate report: '+argdata['reportname'])
        exit(0)

    print('located report '+argdata['reportname'])

    #query recursively to find all child nodes under report node
    querystring = '?uid={}&field={}&op={}&val={}&depth={}&type={}&body=true'.format(\
        reportRootNode[PROP_UID],\
        PROP_LASTMOD,\
        'gt',\
        0,\
        20,\
        '')
    print(COLLABLIO_HOST+"/nodes"+querystring)

    jsonResponse =  fetchNodes(querystring)

    for nodeResult in jsonResponse['nodes']:
        storeNode(nodeResult)
        print(str(nodeResult))


    newdoc = docx.Document()
    initialiseDocStyles(newdoc)
    
    sectionDepth = 1

    for section in reportRootNode[PROP_CHILDLIST]:
        generate(newdoc, getNodeForUID(section[PROP_UID]), sectionDepth)
    
    reportSaveName = 'report-'+str(uuid.uuid4()).replace('-','')[:14]+'.docx'
    newdoc.save(reportSaveName)
    print(reportSaveName)
        
except Exception as e:
    print('an exception occurred while generating the report: '+str(e))
    traceback.print_exc()
                
exit(0)

'''

def generateReportForReportNode(reportRootNode):
    try:
        #query recursively to find all child nodes under report node
        querystring = '?uid={}&field={}&op={}&val={}&depth={}&type={}&body=true'.format(\
            reportRootNode[PROP_UID],\
            PROP_LASTMOD,\
            'gt',\
            0,\
            20,\
            '')
        print(COLLABLIO_HOST+"/nodes"+querystring)

        jsonResponse =  fetchNodes(querystring)

        for nodeResult in jsonResponse['nodes']:
            storeNode(nodeResult)
            print(str(nodeResult))


        newdoc = docx.Document()
        initialiseDocStyles(newdoc)
        
        sectionDepth = 1

        for section in reportRootNode[PROP_CHILDLIST]:
            generate(newdoc, getNodeForUID(section[PROP_UID]), sectionDepth)
        
        reportSaveName = 'report-'+str(uuid.uuid4()).replace('-','')[:14]+'.docx'
        newdoc.save(reportSaveName)
        print(reportSaveName)
        return reportSaveName
            
    except Exception as e:
        print('an exception occurred while generating the report: '+str(e))
        traceback.print_exc()
        return ''




#query to find projectname data['projectname']

if not DAEMON_LISTENER_MODE:
    querystring = '?uid=&field={}&op={}&val={}&depth={}&type={}'.format(\
        PROP_LABEL,\
        'eq',\
        urllib.parse.quote(argdata['projectname']),\
        20,\
        TYPE_PROJECT)
    print(COLLABLIO_HOST+"/nodes"+querystring)

    try:
        jsonResponse =  fetchNodes(querystring)
        projectUID = ''
        for nodeResult in jsonResponse['nodes']:
            if nodeResult[PROP_LABEL] == argdata['projectname']:
                projectUID = nodeResult[PROP_UID]
                break
        print('located project '+projectUID)

        #query recursively to find the specified reportname under projectname
        querystring = '?uid={}&field={}&op={}&val={}&depth={}&type={}'.format(\
            projectUID,\
            PROP_LABEL,\
            'eq',\
            urllib.parse.quote(argdata['reportname']),\
            20,\
            TYPE_REPORT)
        print(COLLABLIO_HOST+"/nodes"+querystring)

        jsonResponse =  fetchNodes(querystring)

        reportRootNode = None
        for nodeResult in jsonResponse['nodes']:
            if nodeResult[PROP_LABEL] == argdata['reportname']:
                reportRootNode = nodeResult

        if not reportRootNode:
            print('unable to locate report: '+argdata['reportname'])
            exit(0)

        print('located report '+argdata['reportname'])

        generateReportForReportNode(reportRootNode)

    except Exception as e:
        print('an exception occurred while generating the report: '+str(e))
        traceback.print_exc()
    
    exit(0)





#query to find all new reports since last check
# todo: save the last check time in persistent storage to avoid fetching every single report upon the process running

lastFetchTime = 0

while True:
    querystring = '?uid=&field={}&op={}&val={}&depth={}&type={}'.format(\
        PROP_LASTMOD,\
        'gt',\
        '__LASTMODTIME__',\
        20,\
        TYPE_REPORT)
    print('*'*30)
    time.sleep(10)
    querystring = querystring.replace('__LASTMODTIME__',str(lastFetchTime))
    print(COLLABLIO_HOST+"/nodes"+querystring)    
    try:
        jsonResponse =  fetchNodes(querystring)
        print(json.dumps(jsonResponse))
        
        if not 'nodes' in jsonResponse:
            continue

        if 'timestamp' in jsonResponse:
            lastFetchTime = int(jsonResponse['timestamp'])
            print('(cur)lastfetchtime='+str(lastFetchTime)+', str(int(jsonResponse[timestamp])='+str(int(jsonResponse['timestamp'])))

        clearNodeIndex()
        
        uidsOfReports = []
        for nodeResult in jsonResponse['nodes']:
            if nodeResult[PROP_TYPE] == TYPE_REPORT:
                uidsOfReports.append(nodeResult[PROP_UID])
                storeNode(nodeResult)

        # now query the database for report generation jobrequests pending for any of those report nodes 
        jsonResponse = fetchNodesPost(uids = uidsOfReports, typ = TYPE_JOBREQ)

        for nodeResult in jsonResponse['nodes']:
            print(str(nodeResult))
            if (PROP_PARENTLIST in nodeResult) and (len(nodeResult[PROP_PARENTLIST]) > 0):
                jobReqParentUID = nodeResult[PROP_PARENTLIST][0][PROP_UID]
                reportNode = getNodeForUID(jobReqParentUID)
                if reportNode:
                    reportfile = generateReportForReportNode(reportNode)
                    if reportfile:
                        print('about to upload '+reportfile)
                        #apparently python3 urllib doesn't have builtin support for multipart/form-data
                        # there's an implementation here https://pymotw.com/3/urllib.request/

                        #data.append('filedata', input.files[0]);
                        #data.append('type', 'file_upload');
                        #data.append('_p', JSON.stringify(params));
                        #await fetch('upload', {
                        #	params.parentid = node[PROP_PARENTLIST][0].uid;
                        
                        params = { 'parentid': jobReqParentUID }

                        # Create the form with simple fields
                        form = MultiPartForm()
                        form.add_field('type', 'file_upload')
                        form.add_field('_p', json.dumps(params))

                        # Add a fake file
                        form.add_file('filedata', reportfile, fileHandle=open(reportfile, "rb"))

                        # Build the request, including the byte-string
                        # for the data to be posted.
                        data = bytes(form)
                        r = urllib.request.Request(COLLABLIO_HOST+'/upload', data=data) #  'http://127.0.0.1:9123'

                        r.add_header('Content-type', form.get_content_type())
                        r.add_header('Content-length', len(data))

                        print()
                        print('OUTGOING DATA:')
                        for name, value in r.header_items():
                            print('{}: {}'.format(name, value))

                        #print(r.data.decode('utf-8'))

                        respStr = urllib.request.urlopen(r).read().decode('utf-8')
                        print('SERVER RESPONSE:')
                        print(respStr)

                        #move the jobrequest node to the recycle bin
                        rbFolderUID = getRecycleBinFolderUID()
                        print('Recycle Bin UID: '+rbFolderUID)
                        

                        moveData = { 'nodes': [nodeResult[PROP_UID]], 'parents': [jobReqParentUID], 'children': [], 'newparent': rbFolderUID }
                        responseForMove = moveNodesPost(moveData)
                        if responseForMove and 'error' in responseForMove:
                            print('error moving jobrequest to recyclebin '+str(responseForMove))
                        else:
                            #delete the local report file
                            os.remove(reportfile)
            
    except Exception as e:
        print('an exception occurred while fetching the report/job nodes: '+str(e))
        traceback.print_exc()


                
exit(0)



'''
TESTJSON = "{\"ops\":[{\"insert\":\"testaetasvsdvsdvdsvv\\n\\nlvjal vlk\\n\"},{\"attributes\":{\"bold\":true},\"insert\":\"dsvklsdvkjdsvjlksdv\"},{\"insert\":\"\\ndv\\n\\nsd\"},{\"attributes\":{\"header\":1},\"insert\":\"\\n\"},{\"insert\":\"vsdvsdsvsdvsdvdv\"},{\"attributes\":{\"list\":\"bullet\"},\"insert\":\"\\n\"},{\"insert\":\"byiuib\"},{\"attributes\":{\"list\":\"bullet\"},\"insert\":\"\\n\"},{\"insert\":\"skjhlkfh\"},{\"attributes\":{\"list\":\"bullet\"},\"insert\":\"\\n\"},{\"insert\":\"fhsdfhfh\"},{\"attributes\":{\"list\":\"bullet\"},\"insert\":\"\\n\"},{\"insert\":\"\\nXML is an inherently \"},{\"attributes\":{\"italic\":true},\"insert\":\"hierarchical \"},{\"insert\":\"data format, and the most natural way to represent it is with a tree. ET has two classes for this purpose -\\n\\n ElementTree represents the whole XML document as a tree, and Element represents a single node in this tree. Interactions with the whole \"},{\"attributes\":{\"code-block\":true},\"insert\":\"\\n\"},{\"insert\":\"\\ndocument (reading and writing to/from files) are usually done on the ElementTree level. Interactions with a single XML element and its sub-elements are done on the Element level.\"},{\"attributes\":{\"blockquote\":true},\"insert\":\"\\n\"},{\"insert\":\"\\nsdvsadvasdvsdvsdvasevlajs v980u2 hsjdkvlsjdv\\n\\n\\nasclnkcj sdchdksjhkjdhjvkdhalsje;f lsej lfsjaelkfjksjdlkjds\"},{\"attributes\":{\"code-block\":true},\"insert\":\"\\n\"},{\"insert\":\"dljsldkjfsjdlfjskdfj;sldj flskdjfklsdjf\\n\\nflasjdfjslkdjkjdvsdlvjlsdkjv dlsvj\"},{\"attributes\":{\"blockquote\":true},\"insert\":\"\\n\"},{\"insert\":\"\\ndlksdjvsdj\\n\\n\\nhat works perfectly... but I don't fully understand what's going on. – some1 Oct 23 '11 at 12:43\\n4\\n\\n, where each element in that list is e + d\"},{\"attributes\":{\"code-block\":true},\"insert\":\"\\n\"},{\"insert\":\", where each element in that list is e + d\"},{\"attributes\":{\"code-block\":true},\"insert\":\"\\n\"},{\"insert\":\", where each element in that list is e + d\"},{\"attributes\":{\"code-block\":true},\"insert\":\"\\n\"},{\"insert\":\"\\n@some1 it basically iterates over the results of the split and adds the delimiter back in. \\\"s is a list, where each element in that list is e + d, where e are the elements in the result of line.split(d), but only if e isn't empty\\\" – JHixson Jun 26 '14 at 17:04\\n\\n\"},{\"attributes\":{\"bold\":true},\"insert\":\"test me out blod \"},{\"insert\":\"then now what?\\nlskdfj\\n\\nlistoof stuff\\nasdkfjldj\"},{\"attributes\":{\"list\":\"bullet\"},\"insert\":\"\\n\"},{\"insert\":\"tesltjk csdlkjkl\"},{\"attributes\":{\"indent\":1,\"list\":\"bullet\"},\"insert\":\"\\n\"},{\"insert\":\"alsdkjfl\"},{\"attributes\":{\"list\":\"bullet\"},\"insert\":\"\\n\"},{\"insert\":\"dfjalskdjf\"},{\"attributes\":{\"list\":\"bullet\"},\"insert\":\"\\n\"},{\"insert\":\"sdjfl\"},{\"attributes\":{\"list\":\"bullet\"},\"insert\":\"\\n\"},{\"insert\":\"\\n\\nffafe\\n\\n\\n\\n\\n\\n\\n\\n\\n\\n\\n\\n\\n\"}]}"

START_HEADING_LEVEL = 0
MAX_HEADING_LEVEL = 9

# https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
def applyStylesToRun(style, run):
    if not style:
        return
    if 'bold' in style:
        run.bold = True
    if 'italic' in style:
        run.italic = True


def initialiseDocStyles(doc):
    if 'Code Block' not in doc.styles:
        codeBlockStyle = doc.styles.add_style('Code Block',docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        codeBlockStyle.base_style = doc.styles['Body Text']
        codeBlockStyle.font.name = 'Consolas'
    
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def applyStylesToParagraph(style, paragraph, doc):
    if not style:
        return
    docStyles = doc.styles
    pformat = paragraph.paragraph_format
    if 'code-block' in style:
        testTable = doc.add_table(1,1)
        testPara = testTable.cell(0,0).paragraphs[0]
        testTable.style = docStyles['Medium Shading 1 Accent 1'] #need to change this to use style Name (not style ID) to get rid of warning
        testTable.style.font.bold = False
        #testTable.style.font.italic = True
        testPara.style.font.bold = False
        testPara.paragraph_format.space_before = Cm(0.1)
        testPara.paragraph_format.line_spacing = 1

        for run in paragraph.runs:
            testPara.add_run(run.text, run.style)

        testPara.style = docStyles['Code Block']
        delete_paragraph(paragraph)
        trailingPara = doc.add_paragraph()
        trailingPara.paragraph_format.space_after = Cm(0)

        return # we've nuked paragraph, nothing else can reference it at this point

    if 'list' in style:
        listItemType = style['list']
        paragraph.style = docStyles['List Bullet']
        pformat.left_indent = Cm(1)
        pformat.space_before = Cm(0)
        pformat.space_after = Cm(0)

    if 'header' in style:
        headingLevel = int(style['header'])
        headingLevel += START_HEADING_LEVEL
        if headingLevel > MAX_HEADING_LEVEL:
            headingLevel = MAX_HEADING_LEVEL
        paragraph.style = docStyles['Heading '+str(headingLevel)]

    if 'blockquote' in style:
        paragraph.style = docStyles['Quote']


    #ensure this is last
    if 'indent' in style:
        pformat.left_indent += Cm(int(style['indent']))
        #run.text = '\t'*int(style['indent'])+run.text

STYLES_WITH_SEPARATE_PARAGRAPHS = ['list']

def isStyleWithItsOwnParagraph(style):
    if not style:
        return False
    for key in list(style.keys()):
        if key in STYLES_WITH_SEPARATE_PARAGRAPHS:
            return True
    #print('key '+key+' is not in STYLES_WITH_SEPARATE_PARAGRAPHS')
    return False


quillOpsData = json.loads(TESTJSON)

# formatting attributes in Quill are generally applied to the last operation's inserted line
# some attributes like bold, italic, strikethrough are applied to the line that follows (forward-applied attributes):
FWD_APPLIED_ATTRIBS = ['bold','italic']

Lines = [{'content':'','formatting':''}]

count = 1
for op in quillOpsData['ops']:
    #print('\n\nLine #{} *********'.format(count))
    count = count+1
    #print(op)
    fwdAttrs = None
    if 'attributes' in op: 
        attrs = op['attributes']
        #print('attrs='+str(attrs))
        opHasFwdAppliedAttrib = False
        #print(list(attrs.keys()))
        for attrName in list(attrs.keys()):
            #print('checking FWD_APPLIED_ATTRIBS for '+attrName)
            opHasFwdAppliedAttrib = (opHasFwdAppliedAttrib or (attrName.lower() in FWD_APPLIED_ATTRIBS))
        if not opHasFwdAppliedAttrib:
            #print('No op attributes are in FWD_APPLY_ATTRIBS'+str(attrs))
            Lines[-1]['formatting'] = attrs
            #print('Lines[-1]: '+str(Lines[-1]))
        else:
            #print('Apply forward attribute'+str(attrs))
            fwdAttrs =  attrs

    startFromStrPos = 0
    while True:
        indexOfNextNewline = op['insert'].find('\n',startFromStrPos)
        if indexOfNextNewline < 0:
            #print('no more newlines found')
            Lines.append({ 'content': op['insert'][startFromStrPos:], 'formatting' : fwdAttrs })
            break
        #print('newline found at: '+str(indexOfNextNewline))
        Lines.append({ 'content': op['insert'][startFromStrPos:indexOfNextNewline+1], 'formatting' : fwdAttrs })
        startFromStrPos = indexOfNextNewline+1
        if startFromStrPos >= len(op['insert']):
            break

cleanedLines = [Lines[0]]
for iL in range(0, len(Lines)):
    line = Lines[iL]
    #print(line)
    #merge lines that are just {line:'\n','formatting':None} with previous line if previous line isn't a sole newline 
    lineBefore = Lines[iL-1]['content']
    if line['content'] == '\n' and not line['formatting'] and lineBefore != '\n':
        cleanedLines[-1]['content'] += '\n'
        #print('dropping empty line')
        continue
    else:
        cleanedLines.append(line)
        
for line2 in cleanedLines:
    print(str(line2['formatting'])+' '+line2['content'])


newdoc = docx.Document()
initialiseDocStyles(newdoc)
newdoc.add_heading('Test Report')
newdoc.add_paragraph('\n')


curLineIndex = 0
prevLineStyle = {}
runStr = ''
prevStyleIsItsOwnParagraph = False

while curLineIndex < len(cleanedLines):

    p = newdoc.add_paragraph('')
    finishedParagraph = False

    while not finishedParagraph:

        curLine = cleanedLines[curLineIndex]['content']
        curLineStyle = cleanedLines[curLineIndex]['formatting']

        prevStyleIsItsOwnParagraph = isStyleWithItsOwnParagraph(prevLineStyle)

        #print('*'*10+'\nprevStyle='+str(prevLineStyle)+' , curStyle='+str(curLineStyle)+', curLine=[['+curLine+']]')

        # the approach is to aggregate lines of the same style into the one paragraph
        # create a new paragraph once the style changes
        
        if ((curLineIndex >= len(cleanedLines) - 1) or ((curLineStyle != prevLineStyle) and runStr) or prevStyleIsItsOwnParagraph):
            #print('  - in add_run area')
            runEndsWithNewline = runStr.endswith('\n')
            if runEndsWithNewline:
                runStr = runStr.strip('\n') #runStr[:-1] #get rid of trailing newline because space is automatically added between paragraphs
            run = p.add_run(runStr)
            
            if prevLineStyle:
                applyStylesToRun(prevLineStyle, run)
                
            #prevStyleIsItsOwnParagraph = isStyleWithItsOwnParagraph(prevLineStyle)
            #if prevStyleIsItsOwnParagraph:
            #    print('* '+str(prevLineStyle)+' is its own paragraph')
            finishedParagraph = (runEndsWithNewline or prevStyleIsItsOwnParagraph or (curLineIndex >= len(cleanedLines) - 1))
            if finishedParagraph:
                applyStylesToParagraph(prevLineStyle, p, newdoc)
            runStr = curLine #commence a new run string, assigning the current line
        else:
            runStr += curLine #append this line to the current run string

        prevLineStyle = curLineStyle
        curLineIndex += 1







reportSaveName = 'report-'+str(uuid.uuid4()).replace('-','')[:14]+'.docx'
newdoc.save(reportSaveName)
'''

'''
############################################################
##### this is the closest to working version so far #######
curindex = 0
while curindex < len(Lines2):
    p = newdoc.add_paragraph('')
    finished = False
    prevStyle = {}
    while not finished:
        curLine = Lines2[curindex]['line']
        curStyle = Lines2[curindex]['formatting']
        if curLine.endswith('\n') and curStyle != prevStyle:
#        if (curindex >= len(Lines2) - 1) or curStyle != prevStyle:
            curLine = curLine[:-1]
            finished = True
        p.add_run(curLine)
#        print('curIndex='+str(curindex)+' runStr='+runStr)
        prevStyle = curStyle
        curindex += 1
############################################################
'''
'''
# a document template needs to have default styles set otherwise python-docx will crash
#newdoc = docx.Document('ReportTemplate.docx')
newdoc = docx.Document()
newdoc.add_heading('TestHeading')
newdoc.add_paragraph('\n')
newdoc.add_paragraph('hello docx')
newdoc.add_page_break()
newdoc.add_paragraph('hello again docx :/')
reportSaveName = 'report-'+str(uuid.uuid4())+'.docx'
newdoc.save(reportSaveName)
'''


'''
curindex = 0
while curindex < len(Lines2):
    p = new paragraph
    finished = False
    while not finsihed:
        p.add_run(Lines2[curindex], <style based on formatting vals>)
        if Lines2[curindex].endswith('\n'):
            finished = true
        curindex += 1
        
        
        
Normal
Header
Header Char
Footer
Footer Char
Heading 1
Heading 2
Heading 3
Heading 4
Heading 5
Heading 6
Heading 7
Heading 8
Heading 9
Default Paragraph Font
Normal Table
No List
No Spacing
Heading 1 Char
Heading 2 Char
Heading 3 Char
Title
Title Char
Subtitle
Subtitle Char
List Paragraph
Body Text
Body Text Char
Body Text 2
Body Text 2 Char
Body Text 3
Body Text 3 Char
List
List 2
List 3
List Bullet
List Bullet 2
List Bullet 3
List Number
List Number 2
List Number 3
List Continue
List Continue 2
List Continue 3
macro
Macro Text Char
Quote
Quote Char
Heading 4 Char
Heading 5 Char
Heading 6 Char
Heading 7 Char
Heading 8 Char
Heading 9 Char
Caption
Strong
Emphasis
Intense Quote
Intense Quote Char
Subtle Emphasis
Intense Emphasis
Subtle Reference
Intense Reference
Book Title
TOC Heading
Table Grid
Light Shading
Light Shading Accent 1
Light Shading Accent 2
Light Shading Accent 3
Light Shading Accent 4
Light Shading Accent 5
Light Shading Accent 6
Light List
Light List Accent 1
Light List Accent 2
Light List Accent 3
Light List Accent 4
Light List Accent 5
Light List Accent 6
Light Grid
Light Grid Accent 1
Light Grid Accent 2
Light Grid Accent 3
Light Grid Accent 4
Light Grid Accent 5
Light Grid Accent 6
Medium Shading 1
Medium Shading 1 Accent 1
Medium Shading 1 Accent 2
Medium Shading 1 Accent 3
Medium Shading 1 Accent 4
Medium Shading 1 Accent 5
Medium Shading 1 Accent 6
Medium Shading 2
Medium Shading 2 Accent 1
Medium Shading 2 Accent 2
Medium Shading 2 Accent 3
Medium Shading 2 Accent 4
Medium Shading 2 Accent 5
Medium Shading 2 Accent 6
Medium List 1
Medium List 1 Accent 1
Medium List 1 Accent 2
Medium List 1 Accent 3
Medium List 1 Accent 4
Medium List 1 Accent 5
Medium List 1 Accent 6
Medium List 2
Medium List 2 Accent 1
Medium List 2 Accent 2
Medium List 2 Accent 3
Medium List 2 Accent 4
Medium List 2 Accent 5
Medium List 2 Accent 6
Medium Grid 1
Medium Grid 1 Accent 1
Medium Grid 1 Accent 2
Medium Grid 1 Accent 3
Medium Grid 1 Accent 4
Medium Grid 1 Accent 5
Medium Grid 1 Accent 6
Medium Grid 2
Medium Grid 2 Accent 1
Medium Grid 2 Accent 2
Medium Grid 2 Accent 3
Medium Grid 2 Accent 4
Medium Grid 2 Accent 5
Medium Grid 2 Accent 6
Medium Grid 3
Medium Grid 3 Accent 1
Medium Grid 3 Accent 2
Medium Grid 3 Accent 3
Medium Grid 3 Accent 4
Medium Grid 3 Accent 5
Medium Grid 3 Accent 6
Dark List
Dark List Accent 1
Dark List Accent 2
Dark List Accent 3
Dark List Accent 4
Dark List Accent 5
Dark List Accent 6
Colorful Shading
Colorful Shading Accent 1
Colorful Shading Accent 2
Colorful Shading Accent 3
Colorful Shading Accent 4
Colorful Shading Accent 5
Colorful Shading Accent 6
Colorful List
Colorful List Accent 1
Colorful List Accent 2
Colorful List Accent 3
Colorful List Accent 4
Colorful List Accent 5
Colorful List Accent 6
Colorful Grid
Colorful Grid Accent 1
Colorful Grid Accent 2
Colorful Grid Accent 3
Colorful Grid Accent 4
Colorful Grid Accent 5
Colorful Grid Accent 6
'''


'''
if __name__ == '__main__':
    # Create the form with simple fields
    form = MultiPartForm()
    form.add_field('firstname', 'Doug')
    form.add_field('lastname', 'Hellmann')

    # Add a fake file
    form.add_file(
        'biography', 'bio.txt',
        fileHandle=io.BytesIO(b'Python developer and blogger.'))

    # Build the request, including the byte-string
    # for the data to be posted.
    data = bytes(form)
    r = request.Request('http://localhost:8080/', data=data)
    r.add_header(
        'User-agent',
        'PyMOTW (https://pymotw.com/)',
    )
    r.add_header('Content-type', form.get_content_type())
    r.add_header('Content-length', len(data))

    print()
    print('OUTGOING DATA:')
    for name, value in r.header_items():
        print('{}: {}'.format(name, value))
    print()
    print(r.data.decode('utf-8'))

    print()
    print('SERVER RESPONSE:')
    print(request.urlopen(r).read().decode('utf-8'))
    
'''